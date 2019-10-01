# imports
from django.shortcuts import render, get_object_or_404, redirect, HttpResponse
from courses.forms import *
from django.db.models import Q
import json
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required, user_passes_test, permission_required
from django.contrib.auth.models import Permission
from songs.models import *
from videos.models import *
from docx import Document
import docx
import spotipy
import spotipy.util as util
import datetime
from django.contrib.messages import error, success
from django.contrib import messages
from accounts.models import User
from django.views import View
from django.contrib.auth.decorators import permission_required
from django.utils.decorators import method_decorator
from accounts.models import SpotifyToken

import os
import spotipy.oauth2 as oauth2
from django.conf import settings

# End: imports -----------------------------------------------------------------

# Functions
def trace(x):
    print(json.dumps(x, indent=4, sort_keys=True))

# End: Functions ---------------------------------------------------------------

addCourse_dec = [
    login_required,
    permission_required('courses.add_course', login_url='forbidden')
]
@method_decorator(addCourse_dec, name='dispatch')
class AddCourseView(View):
    template = 'courses/course_form.html'
    courseForm_class = CourseForm
    sectionForm_class = SectionForm

    def get(self, request):
        courseForm = self.courseForm_class()
        sectionForm = self.sectionForm_class()
        return render(request, 'courses/course_form.html', {
            'courseForm': courseForm,
            'sectionForms': [],
            'sectionFormTemplate': self.sectionForm_class(prefix="template"),
        })

    def post(self, request):

        courseForm = self.courseForm_class(data=request.POST)
        sectionCount = int(request.POST.get("sectionCount", "0"))

        if courseForm.is_valid():
            prefixes = request.POST.getlist("prefix")

            sectionForms = [self.sectionForm_class( prefix=prefixes[i], data=request.POST) for i in range(sectionCount)]

            if all(sectionForm.is_valid() for sectionForm in sectionForms):
                course = courseForm.save()
                duration = 0
                # Add current sections
                for i in range(len(sectionForms)):
                    section = sectionForms[i].save()
                    section.nr = i+1
                    section.course = course
                    if course.start:
                        section.start = course.start + datetime.timedelta(minutes=duration)
                        duration += section.varighet
                    section.save()

                return redirect('courses:course_view', courseID=course.id)

        return render(request, 'courses/course_form.html', {
            'courseForm': courseForm,
            'sectionForms': sectionForms,
            'sectionFormTemplate': self.sectionForm_class(prefix="template"),
        })



editCourse_dec = [
    login_required,
    permission_required('courses.change_course', login_url='forbidden')
]
@method_decorator(editCourse_dec, name='dispatch')
class EditCourseView(View):
    template = 'courses/course_form.html'
    courseForm_class = CourseForm
    sectionForm_class = SectionForm

    def get(self, request, courseID):
        course = Course.objects.get(id=courseID)
        sections = list(course.sections.all())
        courseForm = self.courseForm_class(instance=course)
        sectionForms = [self.sectionForm_class(prefix=i+1, instance=sections[i]) for i in range(len(sections))]

        return render(request, self.template, {
            'courseForm': courseForm,
            'sectionForms': sectionForms,
            'courseID': course.id,
            'sectionFormTemplate': self.sectionForm_class(prefix="template"),
        })

    def post(self, request, courseID):
        course = Course.objects.get(id=courseID)
        sections = list(course.sections.all())
        sectionForms = [self.sectionForm_class(prefix=i+1, instance=sections[i]) for i in range(len(sections))]

        courseForm = CourseForm(request.POST, instance=course)
        sectionCount = int(request.POST.get("sectionCount", "0"))

        if courseForm.is_valid():
            prefixes = request.POST.getlist("prefix")

            sectionForms = [SectionForm( prefix=prefixes[i], data=request.POST) for i in range(sectionCount)]

            if all(sectionForm.is_valid() for sectionForm in sectionForms):
                course = courseForm.save()
                course.sections.all().delete() # reset sections

                duration = 0
                # Add current sections
                for i in range(len(sectionForms)):
                    section = sectionForms[i].save()
                    section.nr = i+1
                    section.course = course
                    if course.start:
                        section.start = course.start + datetime.timedelta(minutes=duration)
                        duration += section.varighet
                    section.save()

                return redirect('courses:course_view', courseID=courseID)

        return render(request, self.template, {
            'courseForm': courseForm,
            'courseID': course.id,
            'sectionForms': sectionForms,
            'sectionFormTemplate': self.sectionForm_class(prefix="template"),
        })



allCourses_dec = [
    login_required,
    permission_required('courses.view_course', login_url='forbidden')
]
@method_decorator(allCourses_dec, name='dispatch')
class AllCoursesView(View):
    template = 'courses/all_courses.html'

    def get(self, request):
        courses = Course.objects.all()
        return render(request, self.template, {'courses': courses})


course_dec = [
    login_required,
    permission_required('courses.view_course', login_url='forbidden')
]
@method_decorator(course_dec, name='dispatch')
class CourseView(View):
    template = 'courses/course_view.html'

    def get(self, request, courseID, ):
        print(request.user.has_perm('courses.view_course'))
        course = Course.objects.get(id=courseID)
        return render(request, self.template, {'course': course})


deleteCourse_dec = [
    login_required,
    permission_required('courses.delete_course', login_url='forbidden')
]
@method_decorator(deleteCourse_dec, name='dispatch')
class DeleteCourseView(View):

    def post(self, request, courseID):
        Course.objects.get(id=courseID).delete()
        success(request, 'Course was successfully deleted')
        return redirect('courses:all_courses')



genPlaylist_dec = [
    login_required,
    permission_required('courses.view_course', login_url='forbidden')
]
@method_decorator(genPlaylist_dec, name='dispatch')
class CreatePlaylistView(View):

    def get(self, request, courseID):
        course = Course.objects.get(id=courseID)

        # Auth client
        sp_oauth = oauth2.SpotifyOAuth(settings.SPOTIFY_CLIENT_ID, settings.SPOTIFY_CLIENT_SECRET, settings.SPOTIFY_REDIRECT_URI, scope=settings.SPOTIFY_SCOPE)

        sp_token, created = SpotifyToken.objects.get_or_create(user=request.user)

        token_info = None
        if sp_token.info:
            token_info = json.loads(sp_token.info)

            if sp_oauth._is_token_expired(token_info):

                print(token_info)
                token_info = sp_oauth.refresh_access_token(token_info['refresh_token'])
                print(token_info)
                sp_token.addInfo(token_info)


            # if scopes don't match, then bail
            # if 'scope' not in token_info or not sp_oauth._is_scope_subset(sp_oauth.scope, token_info['scope']):
            #     messages.error(request, "scope mismatch")
            #     return redirect('courses:course_view', courseID=courseID)

            #return self.token_info['access_token']

        if not token_info:
            auth_url = sp_oauth.get_authorize_url()
            try:
                print(auth_url)
                import webbrowser
                webbrowser.open(auth_url+'&show_dialog=true')
            except Exception as e:
                print(e)
                messages.error(e)
            messages.error(request, 'Due to no connection to Spotify, the playlist was not created. Please try again')
            return redirect('courses:course_view', courseID=courseID)


        token = token_info['access_token']

        # Spotify API object
        spotify = spotipy.Spotify(auth=token)

        # Get username
        spotify_username = spotify.current_user()['uri'].split(':')[-1]

        # Title of playlist to be created
        playlist_title = '{} ({})'.format(course.tittel, course.dato)

        # Get existing playlists for users
        playlists = spotify.user_playlists(spotify_username)

        # Check if playlist already exists
        playlist = [ v for v in playlists['items'] if v['name'] == playlist_title ]

        if not playlist:
            # Create new playlist
            playlist = spotify.user_playlist_create(user=spotify_username, name=playlist_title)
        else:
            # Take first mathing playlist (should be only one anyway)
            playlist = playlist[0]

        # Get URI for all songs used in course
        tracks = [section.song.spotify_URI for section in course.sections.all() if section.song]

        # Add tracks to playlist
        spotify.user_playlist_replace_tracks(user=spotify_username, playlist_id=playlist['id'], tracks=tracks)

        success(request, "Playlist was created, and it's lit!")
        return redirect('courses:course_view', courseID=courseID)


export_dec = [
    login_required,
    permission_required('courses.view_course', login_url='forbidden')
]
@method_decorator(export_dec, name='dispatch')
class ExportView(View):

    def get(self, request, courseID):

        course = Course.objects.get(id=courseID)

        document = Document()

        document.add_heading(course.tittel, level=1)

        tag_names = ", ".join(course.getTags())

        if course.fører:
            fører_navn = course.fører.get_full_name()
        else:
            fører_navn = 'Ingen instruktør valgt'

        if course.følger:
            følger_navn = course.følger.get_full_name()
        else:
            følger_navn = 'Ingen instruktør valgt'


        informasjon = "Instruktør (fører): {}\nInstruktør (følger): {}\nDato: {}\nNår: {} - {}\nHvor: {}\nTema: {}".format(
            fører_navn, følger_navn, course.getDato(), course.getStart(), course.getSlutt(), course.sted, tag_names
        )
        p = document.add_paragraph(informasjon)

        for section in course.sections.all():
            # p = document.add_paragraph("")
            h = document.add_heading("{} ({} min) - {}".format(section.tittel, section.varighet, section.getStart()), level=2)

            p = document.add_paragraph()
            run = p.add_run("Sang: {}".format(section.getSong()))
            run.italic = True
            run.font.size = docx.shared.Pt(9)

            run = p.add_run("\n\n{}".format(section.beskrivelse))


        h = document.add_heading("Kommentarer: ", level=2)
        p = document.add_paragraph(course.kommentarer)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename={}.docx'.format(course)
        document.save(response)

        return response
